from docx import Document
from openpyxl import load_workbook
import docx
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert # 워드를 pdf로 변환

# 엑셀에서 값 읽기
load_wb = load_workbook(r'12. Import_Excel_Information_Automatical_Generate_Certification_Of_Completion\수료증명단.xlsx')
load_ws = load_wb.active

# 이름, 생일, 수료증 번호를 리스트로 생성
name_list = []
birthday_list = []
ho_list = []
for i in range(1, load_ws.max_row + 1):
    name_list.append(load_ws.cell(i, 1).value)
    birthday_list.append(load_ws.cell(i, 2).value)
    ho_list.append(load_ws.cell(i, 3).value)

# 각 변수에 저장된 값 출력
print(name_list)
print(birthday_list)
print(ho_list)

# 에게셀에서 읽은 이름 리스트의 길이(=수료증 수)만큼 반복
for i in range(len(name_list)):
    # 지정된 경로의 파일 읽기
    doc = docx.Document(r'12. Import_Excel_Information_Automatical_Generate_Certification_Of_Completion\수료증양식.docx')

    # 폰트, 글씨 크기 설정
    style = doc.styles['Normal']
    style.font.name = "나눔고딕"
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    style.font.size = docx.shared.Pt(12)

    # 수료증 내용
    para = doc.add_paragraph()
    run = para.add_run('\n\n')
    run = para.add_run('        제' + ho_list[i] + '호\n') # 엑셀에서 읽은 번호를 입력
    run.font.name = '나눔고딕'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    style.font.size = docx.shared.Pt(20)

    para = doc.add_paragraph()
    run = para.add_run('\n\n')
    run = para.add_run('수 료 증')
    run.font.name = '나눔고딕'
    run.bold = True
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    style.font.size = docx.shared.Pt(40)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para = doc.add_paragraph()
    run = para.add_run('\n\n')
    run = para.add_run('        성 명: ' + name_list[i] + '\n') # 엑셀에서 읽은 name_list를 입력
    run.font.name = '나눔고딕'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    style.font.size = docx.shared.Pt(20)
    run = para.add_run('        생 년 월 일: ' + birthday_list[i] + '\n') # 엑셀에서 읽은 birthday_list를 입력
    run.font.name = '나눔고딕'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    style.font.size = docx.shared.Pt(20)
    run = para.add_run('        교 육 과 정:  파이썬과 40개의 작품들\n')
    run.font.name = '나눔고딕'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    style.font.size = docx.shared.Pt(20)
    run = para.add_run('        교 육 날 짜:  2021.08.05~2021.09.09\n')
    run.font.name = '나눔고딕'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    style.font.size = docx.shared.Pt(20)

    para = doc.add_paragraph()
    run = para.add_run('\n\n')
    run = para.add_run('        위 사람은 파이썬과 40개의 작품들 교육과정을\n')
    run.font.name = '나눔고딕'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    style.font.size = docx.shared.Pt(20)
    run = para.add_run('        이수하였으므로 이 증서를 수여합니다.\n')
    run.font.name = '나눔고딕'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    style.font.size = docx.shared.Pt(20)

    para = doc.add_paragraph()
    run = para.add_run('2021.09.19')
    run = para.add_run('\n\n')
    run = para.add_run('파이썬교육기관장')
    run.font.name = '나눔고딕'
    run.bold = True
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    style.font.size = docx.shared.Pt(20)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 이름으로 워드파일 생성
    doc.save('12. Import_Excel_Information_Automatical_Generate_Certification_Of_Completion\\ ' + name_list[i] + '.docx')
    
    # 워드 파일을 읽어 pdf로 변환
    convert('12. Import_Excel_Information_Automatical_Generate_Certification_Of_Completion\\ ' + name_list[i] + '.docx',
            '12. Import_Excel_Information_Automatical_Generate_Certification_Of_Completion\\ ' + name_list[i] + '.pdf')