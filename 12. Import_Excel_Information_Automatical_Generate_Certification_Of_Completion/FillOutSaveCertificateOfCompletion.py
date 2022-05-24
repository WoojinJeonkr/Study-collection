import docx
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 지정된 경로의 파일 읽기
doc = docx.Document(r'12. Import_Excel_Information_Automatical_Generate_Certification_Of_Completion\수료증양식.docx')

# 폰트, 글씨 크기 설정
style = doc.styles['Normal']
style.font.name = "나눔고딕"
style.element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
style.font.size = docx.shared.Pt(12)

# 수료증 내용
para = doc.add_paragraph()
run = para.add_run('\n\n')
run = para.add_run('제 2020-0001 호\n')
run.font.name = '나눔고딕'
style.element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
style.font.size = docx.shared.Pt(20)

para = doc.add_paragraph()
run = para.add_run('\n\n')
run = para.add_run('수 료 증')
run.font.name = '나눔고딕'
run.bold = True
style.element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
style.font.size = docx.shared.Pt(40)
para.alignment = WD_ALIGN_PARAGRAPH.CENTER

para = doc.add_paragraph()
run = para.add_run('\n\n')
run = para.add_run('    성 명:  장다인\n')
run.font.name = '나눔고딕'
style.element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
style.font.size = docx.shared.Pt(20)
run = para.add_run('    생 년 월 일:  2017.12.12\n')
run.font.name = '나눔고딕'
style.element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
style.font.size = docx.shared.Pt(20)
run = para.add_run('    교 육 과 정:  파이썬과 40개의 작품들\n')
run.font.name = '나눔고딕'
style.element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
style.font.size = docx.shared.Pt(20)
run = para.add_run('    교 육 날 짜:  2021.08.05~2021.09.09\n')
run.font.name = '나눔고딕'
style.element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
style.font.size = docx.shared.Pt(20)

para = doc.add_paragraph()
run = para.add_run('\n\n')
run = para.add_run('    위 사람은 파이썬과 40개의 작품들 교육과정을\n')
run.font.name = '나눔고딕'
style.element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
style.font.size = docx.shared.Pt(20)
run = para.add_run('    이수하였으므로 이 증서를 수여합니다.\n')
run.font.name = '나눔고딕'
style.element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
style.font.size = docx.shared.Pt(20)

para = doc.add_paragraph()
run = para.add_run('\n\n\n')
run = para.add_run('2021.09.19')
run.font.name = '나눔고딕'
style.element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
style.font.size = docx.shared.Pt(20)
para.alignment = WD_ALIGN_PARAGRAPH.CENTER

para = doc.add_paragraph()
run = para.add_run('\n\n\n')
run = para.add_run('파이썬교육기관장')
run.font.name = '나눔고딕'
run.bold = True
style.element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
style.font.size = docx.shared.Pt(20)
para.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 수료증결과.docx 파일로 저장
doc.save(r'12. Import_Excel_Information_Automatical_Generate_Certification_Of_Completion\수료증결과.docx')